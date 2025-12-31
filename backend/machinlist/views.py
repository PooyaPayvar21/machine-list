from django.shortcuts import render
from rest_framework import viewsets, status
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework_simplejwt.views import TokenObtainPairView, TokenRefreshView
from django.conf import settings
from .models import User, MachineRegistration
from .serializers import UserSerializer, MachineRegistrationSerializer
from .permissions import IsAdminRole
from rest_framework import permissions
import os
from django.http import HttpResponse
from django.middleware.csrf import get_token
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from docx import Document
from io import BytesIO
from .pdf_utils import fill_machine_pdf

# Create your views here.
class CookieTokenObtainPairView(TokenObtainPairView):
    def finalize_response(self, request, response, *args, **kwargs):
        if response.data.get('access'):
            # Ensure CSRF cookie is set and accessible
            csrf_token = get_token(request)
            response.set_cookie(
                'csrftoken',
                csrf_token,
                httponly=False,
                samesite='Lax'
            )
            
            # Add user info to response body
            try:
                user_email = request.data.get('email')
                if user_email:
                    user = User.objects.get(email=user_email)
                    response.data['user'] = {
                        'id': user.id,
                        'email': user.email,
                        'role': user.role,
                        'is_staff': user.is_staff
                    }
            except User.DoesNotExist:
                pass

            cookie_max_age = 3600 * 24 * 14 # 14 days
            response.set_cookie(
                settings.SIMPLE_JWT['AUTH_COOKIE'],
                response.data['access'],
                max_age=cookie_max_age,
                httponly=settings.SIMPLE_JWT['AUTH_COOKIE_HTTP_ONLY'],
                samesite=settings.SIMPLE_JWT['AUTH_COOKIE_SAMESITE'],
                secure=settings.SIMPLE_JWT['AUTH_COOKIE_SECURE']
            )
            response.set_cookie(
                settings.SIMPLE_JWT['AUTH_COOKIE_REFRESH'],
                response.data['refresh'],
                max_age=cookie_max_age,
                httponly=settings.SIMPLE_JWT['AUTH_COOKIE_HTTP_ONLY'],
                samesite=settings.SIMPLE_JWT['AUTH_COOKIE_SAMESITE'],
                secure=settings.SIMPLE_JWT['AUTH_COOKIE_SECURE']
            )
            del response.data['access']
            del response.data['refresh']
        return super().finalize_response(request, response, *args, **kwargs)

class CookieTokenRefreshView(TokenRefreshView):
    def post(self, request, *args, **kwargs):
        refresh_token = request.COOKIES.get(settings.SIMPLE_JWT['AUTH_COOKIE_REFRESH'])
        if refresh_token:
            request.data['refresh'] = refresh_token
        
        response = super().post(request, *args, **kwargs)
        
        if response.data.get('access'):
            cookie_max_age = 3600 * 24 * 14
            response.set_cookie(
                settings.SIMPLE_JWT['AUTH_COOKIE'],
                response.data['access'],
                max_age=cookie_max_age,
                httponly=settings.SIMPLE_JWT['AUTH_COOKIE_HTTP_ONLY'],
                samesite=settings.SIMPLE_JWT['AUTH_COOKIE_SAMESITE'],
                secure=settings.SIMPLE_JWT['AUTH_COOKIE_SECURE']
            )
            del response.data['access']
        return response

class LogoutView(APIView):
    permission_classes = [permissions.AllowAny]
    authentication_classes = []
    
    def post(self, request):
        response = Response({"message": "Logged out successfully"}, status=status.HTTP_200_OK)
        response.delete_cookie(settings.SIMPLE_JWT['AUTH_COOKIE'])
        response.delete_cookie(settings.SIMPLE_JWT['AUTH_COOKIE_REFRESH'])
        return response

class UserViewSet(viewsets.ModelViewSet):
    queryset = User.objects.all()
    permission_classes = [IsAdminRole]
    serializer_class = UserSerializer

class MachineRegistrationViewSet(viewsets.ModelViewSet):
    queryset = MachineRegistration.objects.all()
    serializer_class = MachineRegistrationSerializer

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            permission_classes = [permissions.IsAuthenticated]
        else:
            permission_classes = [permissions.IsAuthenticated]
        return [permission() for permission in permission_classes]

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def export_machine_doc(request, pk):
    try:
        machine = MachineRegistration.objects.get(pk=pk)
    except MachineRegistration.DoesNotExist:
        return Response({"error": "Machine not found"}, status=404)

    # Template path
    # Try multiple paths to be safe
    possible_paths = [
        os.path.join(settings.BASE_DIR.parent, 'Files', '001-فرم شناسنامه ماشین آلات.docx'),
        r"d:\Pooya\Project\machine-list\Files\001-فرم شناسنامه ماشین آلات.docx"
    ]
    
    template_path = None
    for path in possible_paths:
        if os.path.exists(path):
            template_path = path
            break
            
    if not template_path:
        return Response({"error": "Template file not found"}, status=500)

    try:
        doc = Document(template_path)
    except Exception as e:
        return Response({"error": f"Error loading template: {str(e)}"}, status=500)

    # Prepare data replacements
    replacements = {}
    for field in machine._meta.fields:
        key = f"${{{field.name}}}" # e.g. ${machine_code}
        value = getattr(machine, field.name)
        # Handle date objects
        if value is None:
            value = ""
        else:
            value = str(value)
        replacements[key] = value

    # Helper to replace text in runs to preserve formatting better
    # Simple text replacement in paragraphs
    def process_paragraph(paragraph):
        if not paragraph.text:
            return
        
        # Check if any key is present
        text = paragraph.text
        updated = False
        for key, value in replacements.items():
            if key in text:
                text = text.replace(key, value)
                updated = True
        
        if updated:
            paragraph.text = text

    # Iterate paragraphs
    for p in doc.paragraphs:
        process_paragraph(p)

    # Iterate tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p)
                    
    # Save to buffer
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    
    filename = f"Machine_{machine.machine_code}.docx"
    response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def export_machine_pdf(request, pk):
    try:
        machine = MachineRegistration.objects.get(pk=pk)
    except MachineRegistration.DoesNotExist:
        return Response({"error": "Machine not found"}, status=404)

    try:
        # Prepare data for PDF
        data = MachineRegistrationSerializer(machine).data
        
        pdf_stream = fill_machine_pdf(data)
        
        filename = f"Machine_{machine.machine_code}.pdf"
        response = HttpResponse(pdf_stream.read(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response

    except Exception as e:
        print(f"Error generating PDF: {e}")
        return Response({"error": f"Error generating PDF: {str(e)}"}, status=500)
